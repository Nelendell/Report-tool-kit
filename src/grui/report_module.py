from PyQt5.QtWidgets import QMainWindow, QWidget, QVBoxLayout, QPushButton, QFileDialog, QLineEdit, QLabel, QMessageBox, QHBoxLayout
from PyQt5.QtGui import QIcon, QPixmap, QPalette, QBrush, QFont
from wwfl.data_load import load_dataset
from grplt.plotter import create_plot
from cfg import config
from docx import Document
import os
import tempfile

class ReportWindow(QMainWindow):  # Изменили наследование
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Report module")
        self.setGeometry(200, 200, 1000, 700)
        self.setWindowIcon(QIcon(":/icons/report"))
        
        # Центральный виджет
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
    
    def initUI(self):
        layout = QVBoxLayout()
        
        # Заголовок отчета
        layout.addWidget(QLabel("Report Title:"))
        self.edit_title = QLineEdit("Data Analysis Report")
        layout.addWidget(self.edit_title)
        
        # Кнопки управления
        control_layout = QHBoxLayout()
        
        self.btn_load = QPushButton("Load Data")
        self.btn_plot = QPushButton("Generate Plot")
        self.btn_save = QPushButton("Save Report")
        
        control_layout.addWidget(self.btn_load)
        control_layout.addWidget(self.btn_plot)
        control_layout.addWidget(self.btn_save)
        
        layout.addLayout(control_layout)
        
        # Подключение кнопок
        self.btn_load.clicked.connect(self.load_data)
        self.btn_plot.clicked.connect(self.generate_plot)
        self.btn_save.clicked.connect(self.save_report)
        
        self.setLayout(layout)
    
    def load_data(self):
        default_dir = config.get('default_data_dir')
        file_path, _ = QFileDialog.getOpenFileName(
            self, 
            "Open Data File", 
            default_dir, 
            "CSV Files (*.csv);;Excel Files (*.xlsx);;All Files (*)"
        )
        
        if file_path:
            self.current_data = load_dataset(file_path)
    
    def generate_plot(self):
        if self.current_data is not None:
            title = self.edit_title.text() or "Data Plot"
            self.figure = create_plot(self.current_data, title=title)
    
    def save_report(self):
        if not self.figure:
            QMessageBox.warning(self, "Warning", "Generate a plot first!")
            return
        
        # Запрашиваем шаблон
        default_dir = config.get('default_data_dir')
        template_path, _ = QFileDialog.getOpenFileName(
            self, 
            "Select Template", 
            default_dir, 
            "Word Documents (*.docx)"
        )
        
        if not template_path:
            return
        
        # Сохраняем график во временный файл
        temp_img = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        self.figure.savefig(temp_img.name, dpi=300)
        temp_img.close()
        
        try:
            # Создаем отчет
            doc = Document(template_path)
            
            # Вставляем заголовок
            title = self.edit_title.text()
            if title:
                doc.add_heading(title, level=1)
            
            # Вставляем график
            doc.add_paragraph("Data Visualization:")
            doc.add_picture(temp_img.name, width=docx.shared.Inches(6))
            
            # Сохраняем отчет
            default_dir = config.get('default_rpt_dir')
            file_path, _ = QFileDialog.getSaveFileName(
                self, 
                "Save Report", 
                os.path.join(default_dir, "report.docx"), 
                "Word Documents (*.docx)"
            )
            
            if file_path:
                doc.save(file_path)
                QMessageBox.information(self, "Success", f"Report saved to:\n{file_path}")
        
        finally:
            # Удаляем временный файл
            os.unlink(temp_img.name)